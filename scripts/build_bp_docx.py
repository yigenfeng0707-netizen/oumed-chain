# -*- coding: utf-8 -*-
"""瓯医数链商业计划书（决赛版）Word 生成：python-docx 富文本

- 封面 / 自动目录域 / 页眉页脚页码 / 样式化标题 / 表格 / 300dpi 图表 / 实测截图
- 事实口径：docs/superpowers/specs/2026-09-02-deliverables-redesign.md 事实基准
运行：backend/.venv/Scripts/python scripts/build_bp_docx.py
输出：docs/瓯医数链-商业计划书-决赛版.docx
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHARTS = os.path.join(ROOT, "docs", "BP素材", "charts")
SHOTS = os.path.join(ROOT, "docs", "screenshots")
OUT = os.path.join(ROOT, "docs", "瓯医数链-商业计划书-决赛版.docx")

INK = RGBColor(0x1E, 0x29, 0x3B)
BLUE = RGBColor(0x08, 0x91, 0xB2)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "0891B2"
ZEBRA_BG = "F1F5F9"

BODY_FONT = "微软雅黑"
EN_FONT = "Segoe UI"


# ---------- 底层工具 ----------

def set_run_font(run, size=11, bold=False, color=INK, font=BODY_FONT):
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), font)


def add_para(doc, text="", size=11, bold=False, color=INK, align=None,
             space_after=6, space_before=0, line=1.4, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_rich(doc, segments, size=11, align=None, space_after=6, line=1.4,
             indent=None):
    """segments: list[(text, bold, color)]"""
    p = add_para(doc, "", size=size, align=align, space_after=space_after,
                 line=line, indent=indent)
    for text, bold, color in segments:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = add_para(doc, text, size=17, bold=True, color=BLUE,
                     space_before=14, space_after=8)
        # 底部边框线
        pPr = p._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), HEADER_BG)
        bottom.set(qn("w:space"), "2")
        pbdr.append(bottom)
        pPr.append(pbdr)
        p.style = doc.styles["Heading 1"]
        for run in p.runs:
            set_run_font(run, size=17, bold=True, color=BLUE)
    elif level == 2:
        p = add_para(doc, text, size=13.5, bold=True, color=INK,
                     space_before=10, space_after=5)
        p.style = doc.styles["Heading 2"]
        for run in p.runs:
            set_run_font(run, size=13.5, bold=True, color=INK)
    else:
        p = add_para(doc, text, size=12, bold=True, color=BLUE,
                     space_before=8, space_after=4)
        p.style = doc.styles["Heading 3"]
        for run in p.runs:
            set_run_font(run, size=12, bold=True, color=BLUE)
    return p


def add_bullet(doc, segments, size=11):
    """segments: str 或 list[(text, bold, color)]，前缀圆点"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.35
    pf.left_indent = Cm(0.6)
    pf.first_line_indent = Cm(-0.35)
    run = p.add_run("• ")
    set_run_font(run, size=size, bold=True, color=BLUE)
    if isinstance(segments, str):
        run = p.add_run(segments)
        set_run_font(run, size=size)
    else:
        for text, bold, color in segments:
            run = p.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, widths=None, zebra=True, font_size=10.5,
              highlight_last=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=font_size, bold=True, color=WHITE)
        shade_cell(cell, HEADER_BG)
    # 数据行
    for i, row in enumerate(rows):
        is_last_hl = highlight_last and i == len(rows) - 1
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if j > 0
                           or isinstance(val, str) and len(val) < 14
                           else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(str(val))
            set_run_font(run, size=font_size, bold=is_last_hl,
                         color=BLUE if is_last_hl else INK)
            if is_last_hl:
                shade_cell(cell, "E0F2FE")
            elif zebra and i % 2 == 1:
                shade_cell(cell, ZEBRA_BG)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    add_para(doc, "", size=4, space_after=2)
    return t


def add_figure(doc, filename, width_cm, caption):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    cap = add_para(doc, caption, size=9.5, color=GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    return cap


def add_screenshot(doc, filename, width_cm, caption):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    add_para(doc, caption, size=9.5, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def add_page_number_footer(section):
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def _field(instr):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        return fld
    run = footer_p.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    footer_p._p.append(_field("PAGE"))
    run = footer_p.add_run(" 页 / 共 ")
    set_run_font(run, size=9, color=GRAY)
    footer_p._p.append(_field("NUMPAGES"))
    run = footer_p.add_run(" 页")
    set_run_font(run, size=9, color=GRAY)


def add_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（打开 Word 后右键此处 → 更新域，即可生成目录）"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------- 文档构建 ----------

doc = Document()

# 默认样式
style = doc.styles["Normal"]
style.font.name = EN_FONT
style.font.size = Pt(11)
style.font.color.rgb = INK
style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

# 页面设置 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
add_page_number_footer(section)

# 页眉
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header_p.add_run("瓯医数链 OuMedTrust · 商业计划书（决赛版）")
set_run_font(run, size=8.5, color=GRAY)

# ===== 封面 =====
for _ in range(5):
    add_para(doc, "", space_after=8)
add_para(doc, "第二届全球技术创新大赛 · AI+医疗专题赛", size=13, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "赛道二：医疗大模型与数据", size=12, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
add_para(doc, "瓯医数链 OuMedTrust", size=34, bold=True, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para(doc, "医疗数据要素可信流通平台", size=19, bold=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_para(doc, "让医疗数据「可用不可见、可控可计量」", size=13, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
add_figure(doc, "03_architecture.png", 14.5, "四层架构：治理 · 协作 · 流通 · 监管完整闭环")
add_para(doc, "", space_after=30)
add_para(doc, "2026 年 9 月", size=12, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
page_break(doc)

# ===== 目录 =====
add_heading(doc, "目录", 1)
add_toc(doc)
page_break(doc)

# ===== 一、执行摘要 =====
add_heading(doc, "一、执行摘要", 1)
add_rich(doc, [
    ("瓯医数链是一个", False, INK),
    ("医疗数据要素可信流通平台", True, INK),
    ("：以联邦学习实现医院间「数据不出院」联合建模，以本地大模型实现病历自动治理与 "
     "PHI 脱敏，以审计存证链保障全程合规，并打通数据产品目录、授权交易、收益分成与"
     "监管看板的完整闭环。", False, INK),
], indent=0.75)
add_rich(doc, [
    ("一句话定位：医疗数据界的「可信数据交易所 + 治理工厂」", True, BLUE),
    ("——让沉睡在医院服务器里的数据，变成可确权、可定价、可流通、可监管的数据要素。",
     False, INK),
], indent=0.75)
add_bullet(doc, [("效果已证实：", True, INK),
                 ("297 例真实患者数据上，联邦模型 AUC 0.9091，追平数据大池化集中训练上界 0.9019（合成三院双基准 0.7018 追平 0.7012）", False, INK)])
add_bullet(doc, [("系统已在真实运行：", True, INK),
                 ("魔搭线上 Demo 全时可用；支付宝 live 真实收款闭环 2026-09-02 终验通过——真实订单 ¥3.90 实付落账、交易号回写、存证哈希上链，全程 4 分 42 秒", False, INK)])
add_bullet(doc, [("工程质量可验证：", True, INK),
                 ("448 项单元测试；LLM 全链路并发压测 35/35 请求 100% 成功（含 429 限频主备自动切换实战）", False, INK)])
add_bullet(doc, [("模式已闭环：", True, INK),
                 ("医院 70% / 平台 20% / 数据贡献者 10% 自动分成，监管方实时可见", False, INK)])
add_para(doc, "", space_after=2)
add_rich(doc, [
    ("产品形态：1 个可信数据底座（联邦协作 × AI 治理 × 流通交易 × 监管合规）+ N 个医疗 "
     "AI 应用生态——9 个医疗智能体作为数据消费方与民生出口，其中泛癌卫士直接承载温附医"
     "发表于 Cell（2026）的 Oncoformer 泛癌预测模型，是「医院数据不出院、模型作为数据"
     "产品在平台流通」的标杆案例。", False, INK),
], indent=0.75)
page_break(doc)

# ===== 二、痛点与政策机遇 =====
add_heading(doc, "二、痛点与政策机遇", 1)
add_heading(doc, "2.1 三重痛点", 2)
add_bullet(doc, [("医院不敢共享：", True, INK),
                 ("医疗数据敏感、权属模糊，传统「数据大池化」模式面临《数据安全法》"
                  "《个人信息保护法》硬约束——数据不敢动，AI 模型吃不到跨院数据", False, INK)])
add_bullet(doc, [("数据不会增值：", True, INK),
                 ("基层医院数据质量参差（特征缺失 12%-32%），单院建模 AUC 仅 0.69，"
                  "数据资产无法变现，医院缺乏参与动力", False, INK)])
add_bullet(doc, [("流通没有基础设施：", True, INK),
                 ("数据产品如何定价、如何授权、如何分成、如何监管——缺乏「可用不可见」"
                  "的技术底座与合规流程", False, INK)])
add_heading(doc, "2.2 政策机遇（精准卡位）", 2)
add_bullet(doc, "「数据要素×医疗健康」为国家数据局重点行动方向；浙江省是数据要素市场化配置改革试点")
add_bullet(doc, "本赛道主办方第一名为温州市经济和信息化局——数据要素正是其核心 KPI")
add_bullet(doc, "中国（温州）智能谷定位数字经济，医疗数据要素平台是天然的落地载体")
page_break(doc)

# ===== 三、解决方案 =====
add_heading(doc, "三、解决方案：四层架构闭环", 1)
add_figure(doc, "03_architecture.png", 16.0, "图 3-1  瓯医数链四层架构")
add_heading(doc, "3.1 闭环故事（一条患者数据的旅程）", 2)
add_para(doc, "社区卫生中心的病历经 AI 治理 Copilot 脱敏结构化 → 上架数据要素市场 → "
              "联邦任务消费三院数据训练风险模型（AUC 追平集中训练上界）→ 模型作为数据产品"
              "被健康卫士 Agent 授权采购 → 每一笔交易自动存证上链 → 监管方看板实时可见："
              "流通金额、收益分配、隐私事件 0 起。", indent=0.75)
add_heading(doc, "3.2 为什么评委可以立刻验证", 2)
add_para(doc, "上述闭环不是 PPT 概念：线上 Demo 已部署魔搭全时可访问，数据要素市场已打通"
              "支付宝 live 真实收款（真实订单实付落账），监管看板与管理端对账页实时可见。"
              "详见第六章「已验证的落地实证」。", indent=0.75)
page_break(doc)

# ===== 四、核心技术创新 =====
add_heading(doc, "四、核心技术创新（全部真实现，现场可运行）", 1)
add_heading(doc, "4.1 轻量联邦学习引擎（自研，纯 CPU）", 2)
add_para(doc, "FedAvg + 差分隐私（DP-FedAvg 裁剪 + 高斯噪声）+ 联邦统计。实测：三家异构"
              "医院（三甲 4200 例 / 县医院 2400 例 / 社区 1100 例，特征缺失率 5%-22%）联邦后"
              "全局 AUC 0.7018，追平数据大池化集中训练上界 0.7012；在 297 例真实患者数据"
              "（UCI Cleveland）上复验 AUC 0.9091 追平上界 0.9019。逐院公平性全部获益，"
              "基层获益最大。", indent=0.75)
add_heading(doc, "4.2 AI 病历治理 Copilot（本地 qwen3:4b，院内网推理）", 2)
add_para(doc, "非结构化病历 → 结构化 JSON（患者/主诉/诊断/生命体征/用药/既往史 7 字段），"
              "PHI 脱敏规则引擎识别身份证/手机号/姓名/住院号四类实体零漏检；LLM 失效自动"
              "降级规则引擎，系统永不宕机。", indent=0.75)
add_heading(doc, "4.3 审计存证链", 2)
add_para(doc, "每个联邦任务与交易的摘要 sha256 串联上链，任何篡改导致链条断裂，监管方一键"
              "校验；隐私事件实时监控（当前 0 起）。", indent=0.75)
add_heading(doc, "4.4 数据要素交易引擎", 2)
add_para(doc, "产品目录 → 用途限定授权 → 收益分成（医院 70% / 平台 20% / 数据贡献者 10%）"
              "→ 监管统计；已接入支付宝 live 真实收款，沙箱环境分账逻辑 E2E 验证通过。",
         indent=0.75)
add_heading(doc, "4.5 工程底座", 2)
add_bullet(doc, [("448 项单元测试", True, INK),
                 ("（含生产鉴权与越权审计回归），CI/CD，三级降级容错（本地模型 → 云端 API → 规则引擎），Docker Compose 一键部署", False, INK)])
add_bullet(doc, [("LLM 主备双模高可用：", True, INK),
                 ("主力网关 429 限频时自动切换备选模型；推理模型空内容自愈；全链路压测 5 路 / 10 路并发均 100% 成功", False, INK)])
add_figure(doc, "05_loadtest.png", 16.0, "图 4-1  LLM 全链路并发压测（真实 API 调用，2026-09-01）")
page_break(doc)

# ===== 五、实验验证 =====
add_heading(doc, "五、实验验证（种子固定，可复现）", 1)
add_heading(doc, "5.1 双基准 AUC 对比", 2)
add_figure(doc, "01_auc_dual.png", 16.5, "图 5-1  真实患者数据与合成数据双基准：联邦均追平集中训练上界")
add_table(doc,
          ["方案", "UCI 真实数据 AUC", "合成三院 AUC"],
          [
              ["三甲医院本地模型", "—", "0.7013"],
              ["A 机构（老年层）本地模型", "0.8362", "—"],
              ["B 机构（中年层）本地模型", "0.8427", "—"],
              ["C 机构（青年层）本地模型", "0.9351", "—"],
              ["县医院 / 社区中心本地模型", "—", "0.6994 / 0.6896"],
              ["联邦学习 FedAvg（数据不出院）", "0.9091", "0.7018"],
              ["集中训练上界（现实不可行）", "0.9019", "0.7012"],
          ],
          widths=[7.5, 4.3, 4.3], highlight_last=True)
add_heading(doc, "5.2 收敛性与可复现性", 2)
add_figure(doc, "02_convergence.png", 15.5, "图 5-2  UCI 真实患者数据联邦收敛曲线（12 轮稳定收敛）")
add_bullet(doc, "UCI 心脏病真实队列（Cleveland，297 例真实患者、冠脉病变终点）按年龄三分位构造非 IID 三机构场景；数据阳性率贴近真实流行病学")
add_bullet(doc, "合成三院数据阳性率 21.8%-25.1%；AUC 0.70 符合临床文献区间（0.65-0.75）")
add_bullet(doc, "联邦第 2 轮即达 0.9069，12 轮稳定收敛；种子固定，单机 CPU 即可复现")
add_heading(doc, "5.3 治理流水线实测", 2)
add_bullet(doc, "单份病历脱敏 + 结构化约 10-60 秒（本地推理），治理产物零 PHI 残留")
add_bullet(doc, "演示环境：一台普通笔记本（无 GPU）全本地运行，决赛现场零网络依赖")
page_break(doc)

# ===== 六、已验证的落地实证 =====
add_heading(doc, "六、已验证的落地实证（本章全部为真实发生的事实）", 1)
add_heading(doc, "6.1 线上 Demo 全时可访问", 2)
add_para(doc, "平台已部署于魔搭创空间（https://gsym236998-oumed-chain-demo.ms.show），"
              "含 9 智能体生态、泛癌卫士、EEG 真实数据面板与管理后台；管理员强密码 + 密钥"
              "加固已上线；GitHub Actions CI/CD 自动构建部署。", indent=0.75)
add_screenshot(doc, "07-平台首页.png", 15.0, "图 6-1  线上 Demo 平台首页（魔搭创空间）")
add_heading(doc, "6.2 支付宝 live 真实收款闭环（2026-09-02 终验）", 2)
add_para(doc, "这不是沙箱：支付宝开放平台应用已过审上线，生产收银台直达 excashier.alipay.com；"
              "真实订单（¥3.90 健康卫士·深度画像）由第三方扫码实付，后端订单状态 pending → paid，"
              "真实交易号回写，存证哈希生成，管理端对账页可见（mode=live）。下单到落账全程 "
              "4 分 42 秒。", indent=0.75)
add_figure(doc, "04_payment_loop.png", 16.0, "图 6-2  真实收款闭环时间轴（订单 OM260902105109807F）")
add_screenshot(doc, "alipay_live_cashier_checkout.png", 14.0,
               "图 6-3  支付宝生产收银台（live 模式实测）")
add_heading(doc, "6.3 工程可靠性实证", 2)
add_bullet(doc, "并发压测：5 路（验收档）与 10 路（2 倍压力档）全部 100% 成功；主力网关真实 429 限频触发主备自动切换，请求最终 200 返回")
add_bullet(doc, "448 项单元测试通过；健康检查基线 <35ms，延迟几乎全部来自真实模型推理")
add_para(doc, "上述三项实证合在一起回答了评委最关心的问题：这套系统不是原型演示，"
              "而是一个已经在公网上真实运行、能真实收款、能扛住并发的可运营系统。",
         indent=0.75, bold=False)
page_break(doc)

# ===== 七、商业模式 =====
add_heading(doc, "七、商业模式", 1)
add_figure(doc, "06_revenue.png", 16.0, "图 7-1  四大收入源")
add_heading(doc, "7.1 收入模型", 2)
add_table(doc,
          ["收入源", "定价 / 机制", "状态"],
          [
              ["平台佣金", "数据产品交易额的 20%", "分成架构已内置，沙箱分账 E2E 验证"],
              ["模型 API 订阅", "联邦模型按年授权 ¥80,000/年起", "目录已上架，live 收款通道已打通"],
              ["治理服务费", "病历治理与数据资产化服务，按数据集计价", "治理 Copilot 已上线"],
              ["监管 SaaS", "卫健/数据局合规看板年费", "监管看板已实现"],
          ],
          widths=[3.4, 6.6, 6.0])
add_heading(doc, "7.2 目标客户与竞争差异化", 2)
add_bullet(doc, [("目标客户：", True, INK),
                 ("区域医共体（三甲 + 县医院 + 社区中心）、卫健/数据主管部门、保险精算与药研 CRO（数据买方）", False, INK)])
add_bullet(doc, [("差异化：", True, INK),
                 ("不做「又一个医疗大模型」，而是做大模型时代的数据要素基础设施——与医疗 AI 应用是共生关系（我们的应用生态就是数据买方）；与隐私计算厂商相比多了「流通交易 + 监管合规」的完整闭环", False, INK)])
add_figure(doc, "07_sharing.png", 10.5, "图 7-2  收益分成 70/20/10：医院得大头")
page_break(doc)

# ===== 八、温州落地计划 =====
add_heading(doc, "八、温州落地计划", 1)
add_figure(doc, "08_timeline.png", 16.0, "图 8-1  温州落地路线图")
add_table(doc,
          ["阶段", "里程碑"],
          [
              ["落地（T+3 月）", "入驻中国（温州）智能谷；注册项目公司；申请鹿城区「AI+医疗应用场景示范项目库」"],
              ["试点（T+6 月）", "与 1 家三甲 + 2 家基层机构签约，真实医共体环境联邦建模试点"],
              ["展开（T+12 月）", "数据产品目录上线运营，首笔真实数据要素交易；申报省标杆项目（衔接最高 500 万场景应用奖补）"],
              ["规模化（T+24 月）", "复制到温州其他区县及浙南医共体，建成区域医疗数据要素流通基础设施"],
          ],
          widths=[3.6, 12.4])
add_para(doc, "", space_after=2)
add_rich(doc, [
    ("与大赛政策的衔接：", True, INK),
    ("落地后可申报重大科技攻关（最高 200 万）、软件开发（最高 200 万）、"
     "「算力券/模型券/语料券」补助（每年最高 50 万）、领军型创业项目支持（30-3000 万）。",
     False, INK),
], indent=0.75)
page_break(doc)

# ===== 九、团队与知识产权 =====
add_heading(doc, "九、团队与知识产权", 1)
add_heading(doc, "9.1 团队", 2)
add_para(doc, "全栈开发 + 审核文档 + 数据库工程三人核心团队（详见报名表）；AI 辅助开发范式"
              "使工程效率成为竞争壁垒——448 项测试与完整合规文档即为此范式的直接产物。",
         indent=0.75)
add_heading(doc, "9.2 知识产权", 2)
add_table(doc,
          ["类型", "名称 / 登记号", "权属与状态"],
          [
              ["软件著作权（主轨）", "瓯医数链平台 V1.0 / 治理 Copilot / 联邦引擎 ×3", "材料齐备，提交中（受理通知书将作资格凭证）"],
              ["软件著作权（已登记）", "基于目标检测的视频人数统计计算软件 V1.0（2023SR0719236）", "个人独有，2023-06 下证，证书原件已核验；与影像卫士同源"],
              ["软件著作权（已登记）", "基于AI大数据的智能充电桩故障检测计算软件 V1.0（2023SR0718714）", "个人独有，2023-06 下证；时序异常检测，与 EEG/健康预警同源"],
              ["软件著作权（共有）", "商业经营管理数据异常预警软件 V1.0（2023SR0723903）", "与银泰商业管理集团共有；团队技术能力佐证"],
              ["发明专利（在案）", "数据交易的管理方法、装置、电子设备及存储介质（申请号 2025111566552）", "本人为发明人（申请人中国电信），2025-08 申请，与数据要素流通主题高度同源"],
          ],
          widths=[3.4, 7.0, 5.6], font_size=9.5)
add_heading(doc, "9.3 合规资质", 2)
add_para(doc, "全流程对齐《数据安全法》《个人信息保护法》《数据二十条》；对标 2026 年国家"
              "卫健委等 14 部门「严守医疗数据安全」专项要求与 GB/T 39725-2020 数据分类分级"
              "框架（6 大类 × 5 级）——授权矩阵、存证哈希、越权审计均按此设计；去标识化后"
              "降级流通为国标明确的合法通道，支撑「数据不出院、可用不可见」；演示数据全部"
              "为合成数据，零合规风险。已产出《个人信息保护影响评估（PIA）报告》。",
         indent=0.75)
page_break(doc)

# ===== 十、风险与应对 =====
add_heading(doc, "十、风险与应对", 1)
add_table(doc,
          ["风险", "应对"],
          [
              ["医院参与意愿不足", "分成机制医院占 70%，基层获益最大（逐院公平性实测）；治理服务帮医院零成本数据资产化"],
              ["合规监管变化", "架构对齐数据二十条与 GB/T 39725 分类分级；监管方作为平台一方角色内置，政策变化即产品能力"],
              ["模型效果边界", "双基准验证 + 文献区间校准；三级降级容错保证可用性不依赖单一模型"],
              ["单点供应商依赖", "LLM 主备双模已实战验证（429 自动切换）；本地 qwen3:4b 兜底，断网可跑"],
          ],
          widths=[4.6, 11.4])
page_break(doc)

# ===== 十一、现场演示大纲 =====
add_heading(doc, "十一、现场演示大纲（8 分钟动线）", 1)
add_table(doc,
          ["时间", "环节", "内容"],
          [
              ["0:00-1:00", "痛点与定位", "医院数据「不敢共享、不会增值、没有基建」三重困境"],
              ["1:00-2:30", "联邦协作网络", "三家医院数据全景 → 一键联邦训练 → AUC 曲线与逐院公平性 → 追平集中训练上界"],
              ["2:30-4:00", "AI 病历治理", "粘贴真实格式病历 → 一键治理 → 4 类 PHI 实时掩码 → 本地 qwen3:4b 结构化（院内网，断网可跑）"],
              ["4:00-5:30", "数据要素市场", "产品目录 → 一键授权采购 → 支付宝真实收款 → 实时分成与存证哈希"],
              ["5:30-6:30", "监管看板", "流通金额、收益分配、隐私事件 0 起、存证链校验、管理端对账"],
              ["6:30-8:00", "落地与愿景", "温州医共体试点路径 + 政策衔接收尾"],
          ],
          widths=[2.6, 3.4, 10.0])
add_para(doc, "", space_after=6)
add_para(doc, "附件：results/experiment_results.json（可复现实验数据）、docs/真实数据集复验报告.md、"
              "docs/LLM并发压测报告.md、docs/screenshots/（支付宝 live 闭环证据链）、"
              "docs/PIA-个人信息保护影响评估报告.md",
         size=9.5, color=GRAY)

doc.save(OUT)
print("saved:", OUT)
