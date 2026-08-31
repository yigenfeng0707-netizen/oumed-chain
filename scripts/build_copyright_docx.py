# -*- coding: utf-8 -*-
"""软著程序说明书 Markdown -> Word 转换器（python-docx）

规则（中国版权保护中心格式）：
- 每个「## 第N页」页块 = Word 一页（块末插入分页符）
- 正文宋体 12pt（小四），标题黑体，中文字体经 w:eastAsia 显式设置
- Markdown 表格转 Word 表格（表格线 + 表头加粗）
- **bold** / `code` / 列表 / 引用块 基础样式支持
- 1.5 倍行距，使每页文字行数贴近 30 行口径

运行：python scripts/build_copyright_docx.py
输出：docs/软著申请/程序说明书-1/2/3-*.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "docs" / "软著申请"

SPECS = [
    ("程序说明书-1-瓯医数链医疗数据要素可信流通平台V1.0.md",
     "程序说明书-1-瓯医数链医疗数据要素可信流通平台V1.0.docx"),
    ("程序说明书-2-瓯医病历智能治理系统V1.0.md",
     "程序说明书-2-瓯医病历智能治理系统V1.0.docx"),
    ("程序说明书-3-瓯医联邦学习医疗协作引擎V1.0.md",
     "程序说明书-3-瓯医联邦学习医疗协作引擎V1.0.docx"),
]

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman",
                 size=12, bold=False, color=None):
    """同时设置西文与中文字体（eastAsia 必须显式指定，否则中文回退默认）。"""
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)


def add_rich_text(par, text, size=12, bold_all=False):
    """支持 **bold** 与 `code` 的行内富文本。"""
    for seg in INLINE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**") and len(seg) > 4:
            run = par.add_run(seg[2:-2])
            set_run_font(run, "黑体", size=size, bold=True)
        elif seg.startswith("`") and seg.endswith("`") and len(seg) > 2:
            run = par.add_run(seg[1:-1])
            set_run_font(run, east_asia="宋体", ascii_font="Consolas",
                         size=max(size - 1, 10), color=(80, 80, 80))
        else:
            run = par.add_run(seg)
            set_run_font(run, "宋体", size=size, bold=bold_all)


def add_table(doc, rows):
    """Markdown 表格行列表 -> Word 表格。"""
    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        par = cell.paragraphs[0]
        par.text = ""
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(cell_text)
        set_run_font(run, "黑体", size=10.5, bold=True)
    for i, row in enumerate(body):
        for j in range(len(header)):
            cell = table.rows[i + 1].cells[j]
            par = cell.paragraphs[0]
            par.text = ""
            if j < len(row):
                add_rich_text(par, row[j], size=10.5)
    doc.add_paragraph()


def render_lines(doc, lines):
    """页块内逐行渲染。"""
    i = 0
    table_buf = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []
        if not line.strip():
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            i += 1
            continue
        if line.startswith("### "):
            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(6)
            add_rich_text(par, line[4:], size=13, bold_all=True)
        elif line.startswith("## "):
            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(10)
            add_rich_text(par, line[3:], size=15, bold_all=True)
        elif line.startswith("> "):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.5)
            add_rich_text(par, "※ " + line[2:], size=10.5)
        elif line.startswith("    "):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.75)
            run = par.add_run(line)
            set_run_font(run, east_asia="宋体", ascii_font="Consolas", size=10)
        elif line.lstrip().startswith(("- ", "* ")):
            par = doc.add_paragraph(style="List Bullet")
            add_rich_text(par, line.lstrip()[2:])
        elif re.match(r"^\d+\.\s", line):
            par = doc.add_paragraph(style="List Number")
            add_rich_text(par, re.sub(r"^\d+\.\s", "", line))
        elif re.match(r"^[a-j]\.\s", line):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.9)
            add_rich_text(par, line)
        else:
            par = doc.add_paragraph()
            add_rich_text(par, line)
        i += 1
    if table_buf:
        add_table(doc, table_buf)


def convert(md_path, docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5

    raw = md_path.read_text(encoding="utf-8").splitlines()
    title = raw[0].lstrip("# ").strip()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(title)
    set_run_font(run, "黑体", size=16, bold=True)

    blocks, cur = [], None
    for line in raw[1:]:
        if line.startswith("## 第"):
            cur = [line]
            blocks.append(cur)
        elif cur is not None:
            cur.append(line)

    for idx, block in enumerate(blocks):
        lines = [ln for ln in block]
        heading = lines[0][3:].strip()
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(6)
        add_rich_text(par, heading, size=15, bold_all=True)
        render_lines(doc, lines[1:])
        if idx < len(blocks) - 1:
            br = doc.add_paragraph()
            br.add_run().add_break(WD_BREAK.PAGE)

    doc.save(str(docx_path))
    return len(blocks)


def main():
    for md_name, docx_name in SPECS:
        md_path = SRC_DIR / md_name
        if not md_path.exists():
            print("缺源文件: " + md_name)
            continue
        pages = convert(md_path, SRC_DIR / docx_name)
        out = SRC_DIR / docx_name
        print("OK %s: %d 页块, %d KB" % (docx_name, pages, out.stat().st_size // 1024))


if __name__ == "__main__":
    main()